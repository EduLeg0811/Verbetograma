"""Geracao de relatorio DOCX a partir do Markdown formal."""

from __future__ import annotations

from io import BytesIO
import re


def markdown_to_report_docx(markdown: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docx.shared import Pt

    def strip_md(text: str) -> str:
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = text.replace("<br>", "\n")
        return text.strip()

    def add_runs(paragraph, text: str) -> None:
        text = text.replace("<br>", "\n")
        for part in re.split(r"(\*\*.*?\*\*|`[^`]+`)", text):
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
            else:
                paragraph.add_run(part)

    def table_row(line: str) -> list[str]:
        return [strip_md(cell) for cell in line.strip().strip("|").split("|")]

    doc = Document()
    
    # Configure Normal style font properties safely to prevent Word Styles corruption
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)
    
    from docx.oxml.ns import qn
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Arial')
    rFonts.set(qn('w:hAnsi'), 'Arial')

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].startswith("|"):
                block.append(lines[i])
                i += 1
            rows = [table_row(row) for row in block if not re.fullmatch(r"\|?[\s:\-|]+\|?", row.strip())]
            if rows:
                table = doc.add_table(rows=1, cols=len(rows[0]))
                table.style = "Table Grid"
                for idx, value in enumerate(rows[0]):
                    table.rows[0].cells[idx].text = value
                for row in rows[1:]:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row[: len(cells)]):
                        cells[idx].text = value
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            doc.add_heading(strip_md(heading.group(2)), level=level)
            i += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        bullet = re.match(r"^-\s+(.+)$", line)
        if numbered:
            doc.add_paragraph(strip_md(numbered.group(1)), style="List Number")
        elif bullet:
            doc.add_paragraph(strip_md(bullet.group(1)), style="List Bullet")
        else:
            paragraph = doc.add_paragraph()
            for idx, segment in enumerate(line.split("<br>")):
                if idx:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                add_runs(paragraph, segment)
        i += 1

    output = BytesIO()
    doc.save(output)
    return output.getvalue()
